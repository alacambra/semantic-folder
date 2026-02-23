"""Unit tests for description/describer.py — AnthropicDescriber behaviour."""

import base64
from unittest.mock import MagicMock, patch

from anthropic.types import Message, TextBlock, Usage

from semantic_folder.description.describer import (
    DEFAULT_MAX_FILE_CONTENT_BYTES,
    DEFAULT_MAX_RETRIES,
    DEFAULT_REQUEST_DELAY,
    AnthropicDescriber,
    _extract_docx_text,
    _file_extension,
    anthropic_describer_from_config,
)

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _make_describer(
    request_delay: float = 0.0,
) -> tuple[AnthropicDescriber, MagicMock]:
    """Return (describer, mock_anthropic_client).

    Args:
        request_delay: Inter-request delay; defaults to 0.0 for fast tests.
    """
    with patch("semantic_folder.description.describer.anthropic.Anthropic") as mock_cls:
        mock_client = MagicMock()
        mock_cls.return_value = mock_client
        describer = AnthropicDescriber(
            api_key="test-key", model="test-model", request_delay=request_delay
        )
    return describer, mock_client


def _mock_message_response(text: str) -> Message:
    """Create an Anthropic Message response with the given text."""
    return Message(
        id="msg-test",
        type="message",
        role="assistant",
        content=[TextBlock(type="text", text=text)],
        model="test-model",
        stop_reason="end_turn",
        usage=Usage(input_tokens=10, output_tokens=5),
    )


# ---------------------------------------------------------------------------
# __init__ tests
# ---------------------------------------------------------------------------


class TestAnthropicDescriberInit:
    def test_creates_client_with_api_key_and_max_retries(self) -> None:
        with patch("semantic_folder.description.describer.anthropic.Anthropic") as mock_cls:
            AnthropicDescriber(api_key="sk-test-123", model="claude-haiku-4-5-20251001")
            mock_cls.assert_called_once_with(api_key="sk-test-123", max_retries=DEFAULT_MAX_RETRIES)

    def test_creates_client_with_custom_max_retries(self) -> None:
        with patch("semantic_folder.description.describer.anthropic.Anthropic") as mock_cls:
            AnthropicDescriber(api_key="sk-test-123", max_retries=5)
            mock_cls.assert_called_once_with(api_key="sk-test-123", max_retries=5)

    def test_stores_model(self) -> None:
        describer, _ = _make_describer()
        assert describer._model == "test-model"

    def test_stores_request_delay(self) -> None:
        describer, _ = _make_describer(request_delay=2.5)
        assert describer._request_delay == 2.5

    def test_default_request_delay(self) -> None:
        with patch("semantic_folder.description.describer.anthropic.Anthropic"):
            describer = AnthropicDescriber(api_key="test-key")
        assert describer._request_delay == DEFAULT_REQUEST_DELAY


# ---------------------------------------------------------------------------
# _file_extension tests
# ---------------------------------------------------------------------------


class TestFileExtension:
    def test_docx(self) -> None:
        assert _file_extension("report.docx") == ".docx"

    def test_pdf(self) -> None:
        assert _file_extension("invoice.pdf") == ".pdf"

    def test_uppercase(self) -> None:
        assert _file_extension("REPORT.DOCX") == ".docx"

    def test_no_extension(self) -> None:
        assert _file_extension("Makefile") == ""

    def test_multiple_dots(self) -> None:
        assert _file_extension("archive.tar.gz") == ".gz"


# ---------------------------------------------------------------------------
# _extract_docx_text tests
# ---------------------------------------------------------------------------


class TestExtractDocxText:
    def test_extracts_text_from_valid_docx(self) -> None:
        import io

        from docx import Document

        doc = Document()
        doc.add_paragraph("Hello World")
        doc.add_paragraph("Second paragraph")
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        result = _extract_docx_text(docx_bytes)

        assert "Hello World" in result
        assert "Second paragraph" in result

    def test_returns_empty_string_on_invalid_bytes(self) -> None:
        result = _extract_docx_text(b"not a valid docx file")
        assert result == ""


# ---------------------------------------------------------------------------
# summarize_file tests — text path
# ---------------------------------------------------------------------------


class TestSummarizeFileText:
    def test_calls_messages_create_with_correct_params(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("A test summary.")

        describer.summarize_file("report.txt", b"file content here")

        mock_client.messages.create.assert_called_once()
        call_kwargs = mock_client.messages.create.call_args[1]
        assert call_kwargs["model"] == "test-model"
        assert call_kwargs["max_tokens"] == 150
        messages = call_kwargs["messages"]
        assert len(messages) == 1
        assert messages[0]["role"] == "user"
        assert "report.txt" in messages[0]["content"]
        assert "file content here" in messages[0]["content"]

    def test_truncates_content_to_max_bytes(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        large_content = b"x" * (DEFAULT_MAX_FILE_CONTENT_BYTES + 5000)
        describer.summarize_file("big.txt", large_content)

        call_kwargs = mock_client.messages.create.call_args[1]
        prompt_content = call_kwargs["messages"][0]["content"]
        # The prompt should contain at most DEFAULT_MAX_FILE_CONTENT_BYTES worth of 'x'
        assert "x" * DEFAULT_MAX_FILE_CONTENT_BYTES in prompt_content
        assert "x" * (DEFAULT_MAX_FILE_CONTENT_BYTES + 1) not in prompt_content

    def test_handles_binary_content_with_replace(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Binary file.")

        binary_content = bytes(range(256))
        describer.summarize_file("image.bin", binary_content)

        call_kwargs = mock_client.messages.create.call_args[1]
        prompt_content = call_kwargs["messages"][0]["content"]
        # Should not raise — binary decoded with errors="replace"
        assert "image.bin" in prompt_content

    def test_returns_text_from_first_content_block(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response(
            "This is a quarterly report."
        )

        result = describer.summarize_file("report.txt", b"Q1 results...")

        assert result == "This is a quarterly report."

    def test_empty_content_produces_valid_prompt(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Empty file.")

        result = describer.summarize_file("empty.txt", b"")

        assert result == "Empty file."
        mock_client.messages.create.assert_called_once()

    def test_sleeps_before_api_call(self) -> None:
        describer, mock_client = _make_describer(request_delay=0.5)
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        with patch("semantic_folder.description.describer.time.sleep") as mock_sleep:
            describer.summarize_file("report.txt", b"content")

        mock_sleep.assert_called_once_with(0.5)

    def test_no_sleep_when_delay_is_zero(self) -> None:
        describer, mock_client = _make_describer(request_delay=0.0)
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        with patch("semantic_folder.description.describer.time.sleep") as mock_sleep:
            describer.summarize_file("report.txt", b"content")

        mock_sleep.assert_not_called()


# ---------------------------------------------------------------------------
# summarize_file tests — docx path
# ---------------------------------------------------------------------------


class TestSummarizeFileDocx:
    def test_extracts_text_and_sends_as_prompt(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Invoice spec.")

        with patch(
            "semantic_folder.description.describer._extract_docx_text",
            return_value="Invoice flow specification content",
        ):
            result = describer.summarize_file("Invoice_Flow.docx", b"\x50\x4b\x03\x04")

        assert result == "Invoice spec."
        call_kwargs = mock_client.messages.create.call_args[1]
        prompt = call_kwargs["messages"][0]["content"]
        assert "Invoice_Flow.docx" in prompt
        assert "Invoice flow specification content" in prompt

    def test_truncates_extracted_text_to_max_bytes(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        long_text = "y" * (DEFAULT_MAX_FILE_CONTENT_BYTES + 5000)
        with patch(
            "semantic_folder.description.describer._extract_docx_text",
            return_value=long_text,
        ):
            describer.summarize_file("big.docx", b"\x50\x4b")

        call_kwargs = mock_client.messages.create.call_args[1]
        prompt = call_kwargs["messages"][0]["content"]
        assert "y" * DEFAULT_MAX_FILE_CONTENT_BYTES in prompt
        assert "y" * (DEFAULT_MAX_FILE_CONTENT_BYTES + 1) not in prompt

    def test_fallback_when_extraction_fails(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Unknown doc.")

        with patch(
            "semantic_folder.description.describer._extract_docx_text",
            return_value="",
        ):
            describer.summarize_file("broken.docx", b"bad data")

        call_kwargs = mock_client.messages.create.call_args[1]
        prompt = call_kwargs["messages"][0]["content"]
        assert "[could not extract text from broken.docx]" in prompt

    def test_case_insensitive_extension(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        with patch(
            "semantic_folder.description.describer._extract_docx_text",
            return_value="Extracted text",
        ) as mock_extract:
            describer.summarize_file("Report.DOCX", b"\x50\x4b\x03\x04")

        mock_extract.assert_called_once()

    def test_sleeps_before_api_call(self) -> None:
        describer, mock_client = _make_describer(request_delay=0.5)
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        with (
            patch("semantic_folder.description.describer.time.sleep") as mock_sleep,
            patch(
                "semantic_folder.description.describer._extract_docx_text",
                return_value="Extracted text",
            ),
        ):
            describer.summarize_file("doc.docx", b"\x50\x4b\x03\x04")

        mock_sleep.assert_called_once_with(0.5)


# ---------------------------------------------------------------------------
# summarize_file tests — pdf path
# ---------------------------------------------------------------------------


class TestSummarizeFilePdf:
    def test_sends_base64_document_block(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("PDF summary.")

        pdf_bytes = b"%PDF-1.4 fake content"
        result = describer.summarize_file("report.pdf", pdf_bytes)

        assert result == "PDF summary."
        call_kwargs = mock_client.messages.create.call_args[1]
        assert call_kwargs["model"] == "test-model"
        assert call_kwargs["max_tokens"] == 150
        messages = call_kwargs["messages"]
        assert len(messages) == 1
        content_blocks = messages[0]["content"]
        assert isinstance(content_blocks, list)
        assert len(content_blocks) == 2

        # Document block
        doc_block = content_blocks[0]
        assert doc_block["type"] == "document"
        assert doc_block["source"]["type"] == "base64"
        assert doc_block["source"]["media_type"] == "application/pdf"
        expected_b64 = base64.standard_b64encode(pdf_bytes).decode("ascii")
        assert doc_block["source"]["data"] == expected_b64

        # Text block with prompt
        text_block = content_blocks[1]
        assert text_block["type"] == "text"
        assert "report.pdf" in text_block["text"]
        assert "Summarize this file in one sentence" in text_block["text"]

    def test_does_not_truncate_pdf_content(self) -> None:
        """PDF content is sent in full (base64), not truncated to max_file_content_bytes."""
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Large PDF.")

        large_pdf = b"%PDF" + b"\x00" * (DEFAULT_MAX_FILE_CONTENT_BYTES + 5000)
        describer.summarize_file("large.pdf", large_pdf)

        call_kwargs = mock_client.messages.create.call_args[1]
        doc_block = call_kwargs["messages"][0]["content"][0]
        expected_b64 = base64.standard_b64encode(large_pdf).decode("ascii")
        assert doc_block["source"]["data"] == expected_b64

    def test_case_insensitive_pdf_extension(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("Summary.")

        describer.summarize_file("Report.PDF", b"%PDF-1.4")

        call_kwargs = mock_client.messages.create.call_args[1]
        content_blocks = call_kwargs["messages"][0]["content"]
        assert content_blocks[0]["type"] == "document"

    def test_sleeps_before_api_call(self) -> None:
        describer, mock_client = _make_describer(request_delay=0.5)
        mock_client.messages.create.return_value = _mock_message_response("PDF summary.")

        with patch("semantic_folder.description.describer.time.sleep") as mock_sleep:
            describer.summarize_file("report.pdf", b"%PDF-1.4")

        mock_sleep.assert_called_once_with(0.5)


# ---------------------------------------------------------------------------
# classify_folder tests
# ---------------------------------------------------------------------------


class TestClassifyFolder:
    def test_calls_messages_create_with_folder_path_and_files(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("project-docs")

        describer.classify_folder("/drive/root:/Projects", ["readme.md", "plan.docx"])

        mock_client.messages.create.assert_called_once()
        call_kwargs = mock_client.messages.create.call_args[1]
        assert call_kwargs["model"] == "test-model"
        assert call_kwargs["max_tokens"] == 50
        messages = call_kwargs["messages"]
        assert len(messages) == 1
        assert "/drive/root:/Projects" in messages[0]["content"]
        assert "- readme.md" in messages[0]["content"]
        assert "- plan.docx" in messages[0]["content"]

    def test_returns_stripped_text(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("  invoices  \n")

        result = describer.classify_folder("/path", ["invoice.pdf"])

        assert result == "invoices"

    def test_empty_filenames_list(self) -> None:
        describer, mock_client = _make_describer()
        mock_client.messages.create.return_value = _mock_message_response("empty-folder")

        result = describer.classify_folder("/path", [])

        assert result == "empty-folder"
        mock_client.messages.create.assert_called_once()

    def test_sleeps_before_api_call(self) -> None:
        describer, mock_client = _make_describer(request_delay=0.5)
        mock_client.messages.create.return_value = _mock_message_response("project-docs")

        with patch("semantic_folder.description.describer.time.sleep") as mock_sleep:
            describer.classify_folder("/path", ["readme.md"])

        mock_sleep.assert_called_once_with(0.5)


# ---------------------------------------------------------------------------
# anthropic_describer_from_config tests
# ---------------------------------------------------------------------------


class TestAnthropicDescriberFromConfig:
    def test_passes_api_key_and_model_from_config(self) -> None:
        config = MagicMock()
        config.anthropic_api_key = "sk-from-config"
        config.anthropic_model = "claude-haiku-4-5-20251001"
        config.max_file_content_bytes = 16384
        config.anthropic_max_retries = 3
        config.anthropic_request_delay = 1.0

        with patch("semantic_folder.description.describer.anthropic.Anthropic") as mock_cls:
            describer = anthropic_describer_from_config(config)

        mock_cls.assert_called_once_with(api_key="sk-from-config", max_retries=3)
        assert describer._model == "claude-haiku-4-5-20251001"
        assert describer._max_file_content_bytes == 16384

    def test_passes_max_retries_from_config(self) -> None:
        config = MagicMock()
        config.anthropic_api_key = "sk-test"
        config.anthropic_model = "test-model"
        config.max_file_content_bytes = 8192
        config.anthropic_max_retries = 5
        config.anthropic_request_delay = 0.0

        with patch("semantic_folder.description.describer.anthropic.Anthropic") as mock_cls:
            anthropic_describer_from_config(config)

        mock_cls.assert_called_once_with(api_key="sk-test", max_retries=5)

    def test_passes_request_delay_from_config(self) -> None:
        config = MagicMock()
        config.anthropic_api_key = "sk-test"
        config.anthropic_model = "test-model"
        config.max_file_content_bytes = 8192
        config.anthropic_max_retries = 3
        config.anthropic_request_delay = 2.5

        with patch("semantic_folder.description.describer.anthropic.Anthropic"):
            describer = anthropic_describer_from_config(config)

        assert describer._request_delay == 2.5
